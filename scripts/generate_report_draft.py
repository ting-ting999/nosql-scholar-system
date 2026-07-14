from __future__ import annotations

import json
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\zyt18\Downloads\2026春-NoSQL课程报告模板.docx")
OUTPUT = ROOT / "docs" / "NoSQL课程项目报告初稿-基于NoSQL的科研数据分析系统.docx"
SUMMARY = ROOT / "data" / "processed" / "summary.json"


def paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)


def add_block(anchor, lines: list[str]):
    current = anchor
    set_text(anchor, lines[0])
    for line in lines[1:]:
        current = paragraph_after(current, line)
        for run in current.runs:
            run.font.size = Pt(10.5)


def top_list(summary, key, label="name", value="count", limit=5):
    return "；".join(f"{item[label]} {item[value]}" for item in summary[key][:limit])


def report_sections(summary: dict) -> list[list[str]]:
    return [
        [
            "NoSQL 数据库适用于海量、异构、半结构化和高关联数据的组织与分析。本项目以教师提供的 Web of Science 学者科研数据为对象，围绕论文元数据管理、科研统计分析、学术关系图谱构建和语义相似检索展开。论文数据包含标题、作者、摘要、关键词、机构、期刊、年份、被引次数和参考文献等字段，具有明显的半结构化和关系密集特征，因此适合作为 NoSQL 综合应用场景。",
            f"本项目共处理 {summary['total_papers']} 篇论文，识别学者 {summary['total_authors']} 名、关键词 {summary['total_keywords']} 个、期刊或会议来源 {summary['total_journals']} 个。系统综合使用 MongoDB、Neo4j 和 FAISS 思路，实现从数据存储、统计可视化、图谱分析到向量检索的完整流程。",
        ],
        [
            "本项目的主要目标包括：一是完成 Web of Science 文本数据的解析、清洗与标准化；二是使用 MongoDB 作为主存储数据库，完成论文元数据的文档化保存和聚合查询；三是完成年度趋势、学院分布、关键词热点、期刊分布、学科类别和合作机构等多角度可视化分析；四是使用 Neo4j 构建学术知识图谱，并通过 Cypher 查询分析作者合作、学院合作、关键词核心学者和最短合作路径；五是设计并实现基于 FAISS 思路的论文语义向量检索和相似学者发现；六是开发统一 Web 展示系统，便于课程演示和结果说明。",
        ],
        [
            "系统整体流程为：原始 WOS 文本数据 -> 数据解析与清洗 -> 结构化 JSON/CSV 输出 -> MongoDB 文档存储 -> 统计聚合与图表展示 -> Neo4j 图谱 CSV 导入 -> Cypher 关系分析 -> 论文文本向量化 -> 相似检索与推荐 -> Web 前端统一展示。",
            "工具分工如下：MongoDB 负责论文元数据的主存储和常规查询；Neo4j 负责作者、论文、机构、学院、关键词、期刊之间的关系建模和路径分析；FAISS 或向量相似度模块负责论文主题语义检索、相似学者发现和潜在合作推荐；FastAPI 提供后端接口；ECharts 提供统计图表和图谱预览。",
        ],
        [
            "原始数据文件为 scholar-data-full.txt，格式为 Web of Science 纯文本记录，每篇论文以 ER 作为结束标记。主要字段包括 AU/AF 作者、TI 标题、SO 来源期刊或会议、DE 作者关键词、ID Keywords Plus、AB 摘要、C1 地址、C3 机构、CR 参考文献、PY 年份、TC 被引次数、DI DOI、WC 学科类别、SC 研究方向、UT WOS 唯一编号。",
            "预处理方法采用“结构化字段优先 + 规则清洗”的策略。程序识别两位字段码，将缩进行合并到所属字段；对作者、机构、参考文献等多值字段保留列表结构；对关键词和学科类别按分号切分；根据 UT 或 DOI 构造论文唯一编号；从地址字段中抽取学院简称并映射为中文学院名称；对缺失年份、被引次数、关键词等字段进行默认值处理。",
            f"数据清洗后得到 {summary['total_papers']} 条有效论文记录。学院分布中排名靠前的包括：{top_list(summary, 'school_distribution')}。",
        ],
        [
            "本项目选择 MongoDB 作为 Hadoop、Redis、MongoDB 三类技术中的主存储方案。选择理由是 Web of Science 论文元数据字段较多，不同论文字段缺失情况不完全一致，且作者、关键词、机构、参考文献等字段天然适合数组结构保存。MongoDB 的文档模型可以较好地保留原始数据层次，同时支持聚合统计和全文索引。",
            "MongoDB 数据库命名为 nosql_scholar，核心集合为 papers。每篇论文对应一个文档，包含 paper_id、title、abstract、authors、keywords、journal、year、times_cited、school、institutions、references、wos_categories 等字段。系统为 paper_id 创建唯一索引，为 year、school、authors、keywords 创建普通索引，并为 title、abstract、keywords 创建全文索引。",
            "典型聚合查询包括：按 year 分组统计年度发文量和被引次数；按 school 分组统计学院发文量；展开 keywords 字段统计高频关键词；按 journal 分组统计主要期刊来源。这些查询结果既可用于后端分析，也可直接驱动前端图表。",
        ],
        [
            "系统完成了六类可视化图表分析。第一类是年度发文趋势和年度被引趋势，用于观察科研产出随时间的变化。数据中近年论文数量较多，2024 年和 2025 年发文量排名靠前，说明数据集中收录了大量近年成果。",
            f"第二类是学院发文量分析，排名靠前的学院包括：{top_list(summary, 'school_distribution', limit=8)}。第三类是关键词热点分析，高频关键词包括：{top_list(summary, 'keyword_distribution', limit=10)}。这些词反映出人工智能、优化、特征提取、材料性能等方向较活跃。",
            f"第四类是主要期刊和会议来源分析，排名靠前的来源包括：{top_list(summary, 'journal_distribution', limit=8)}。第五类是 WOS 学科类别分析，排名靠前的类别包括：{top_list(summary, 'category_distribution', limit=8)}。第六类是合作机构关系分析，用于观察校内外机构合作情况。",
        ],
        [
            "Neo4j 图谱部分构建六类实体：Author 表示学者，Paper 表示论文，Institution 表示机构，School 表示学院，Keyword 表示关键词，Journal 表示期刊或会议来源。关系类型包括 Author-[:WROTE]->Paper、Author-[:AFFILIATED_WITH]->Institution、Author-[:BELONGS_TO]->School、Paper-[:HAS_KEYWORD]->Keyword、Paper-[:PUBLISHED_IN]->Journal，以及 Author-[:COAUTHORED]-Author。",
            "三元组构建规则为：根据作者列表生成作者撰写论文关系；根据 C3 和 C1 字段生成作者所属机构关系；根据地址字段抽取学院并生成作者隶属学院关系；根据关键词字段生成论文包含关键词关系；根据 SO 字段生成论文发表来源关系；同一篇论文内任意两位作者生成合作作者关系。",
            "项目生成的 Neo4j CSV 文件位于 data/neo4j，导入脚本为 docs/neo4j_import.cypher。代表性查询包括某学者的合作作者及合作次数、学院之间的合作关系、关键词相关核心学者、两位学者之间最短合作路径和机构连接地位分析。",
        ],
        [
            "实体提取采用结构化字段优先方法。作者实体主要来自 AF 字段，缩写作者来自 AU 字段；论文实体由 UT、DOI 或序号生成唯一标识；关键词来自 DE 字段和 ID 字段；期刊来自 SO 字段；机构来自 C3 字段，并结合 C1 地址字段补充；学院主要从 C1 地址中的 Guangdong Univ Technol 后续机构短语抽取。",
            "为减少同名异写和缩写不一致问题，项目对空格、标点和大小写进行了统一处理，并建立了学院英文简称到中文学院名称的映射规则。例如 Sch Automat、Fac Automat、Coll Automat 统一映射为自动化学院；Sch Mat & Energy 映射为材料与能源学院；Sch Electromech Engn 映射为机电工程学院。对于无法识别的学院，保留原始短语或标记为未知学院，便于后续人工核验。",
        ],
        [
            "向量化模块选择论文标题、摘要、作者关键词和 Keywords Plus 作为论文语义文本，并将这些字段拼接后转换为高维文本向量。向量构建后进行 L2 归一化，使内积相似度可以近似表示余弦相似度。系统优先检测 FAISS 环境，若安装 faiss-cpu，则使用 IndexFlatIP 构建向量索引；若本机暂未安装 FAISS，则使用 scikit-learn 余弦相似度作为可运行兜底，保证演示系统稳定运行。",
            "本项目完成了相似论文检索和相似学者发现两类任务。相似论文检索支持输入研究主题，例如 deep learning optimization，返回语义最接近的论文列表及相似度。相似学者发现将作者近年论文标题、摘要和关键词拼接为作者画像文本，再计算作者画像之间的相似度。进一步结合 Neo4j 合作关系，可以筛选出研究方向相似但尚未合作的学者，作为潜在合作推荐。",
        ],
        [
            "系统实现过程包括数据处理、索引构建、后端接口、前端展示和文档整理五部分。数据处理脚本 scripts/prepare_data.py 负责将原始 WOS 文本转换为 JSON、CSV 和 Neo4j 导入文件；scripts/import_mongodb.py 负责将论文文档导入 MongoDB；scripts/build_vector_index.py 负责构建论文向量索引；backend/main.py 提供 FastAPI 接口；frontend/index.html、frontend/app.js 和 frontend/styles.css 提供 Web 展示。",
            "Web 系统包含五个页面：总览页面展示论文、学者、关键词、期刊数量和核心趋势图；统计图表页面展示关键词、期刊、学科类别和合作机构图表；学者画像页面展示某学者的论文数、被引数、学术年龄、关键词、合作作者和代表论文；知识图谱页面展示高被引论文子图和 Cypher 查询；向量检索页面展示相似论文和相似学者结果。",
            "本地运行命令为：先执行 python scripts/prepare_data.py --input scholar-data-full.txt，再执行 python scripts/build_vector_index.py，最后运行 uvicorn backend.main:app --host 127.0.0.1 --port 8000，并访问 http://127.0.0.1:8000。",
        ],
        [
            "通过本次课程项目，我进一步理解了 NoSQL 数据库在科研数据场景中的不同分工。MongoDB 适合保存字段复杂、结构不完全统一的论文元数据；Neo4j 适合表达作者、论文、机构、关键词之间的高关联关系，并支持路径分析和网络结构分析；FAISS 适合处理文本向量相似检索，能够从语义层面发现相似论文和相似学者。",
            "项目中比较困难的部分包括 WOS 多行字段解析、学院名称标准化、图谱三元组构建和向量检索环境兼容。解决方法是优先使用结构化字段，采用规则清洗和映射表降低抽取误差，并将 FAISS 模块设计为可检测、可降级的结构。后续可改进方向包括引入更强的中文和英文学术文本嵌入模型，完善作者消歧算法，进一步结合 Neo4j 图结构和向量相似度做更准确的合作推荐。",
        ],
        [
            "[1] MongoDB Documentation. Aggregation Operations and Indexes.",
            "[2] Neo4j Documentation. Cypher Query Language Manual.",
            "[3] Johnson J, Douze M, Jegou H. Billion-scale similarity search with GPUs. IEEE Transactions on Big Data, 2019.",
            "[4] Clarivate. Web of Science Core Collection Field Tags.",
            "[5] Apache ECharts Documentation. Chart Configuration and Visualization.",
        ],
        [
            "附录文件包括：data/processed/papers.json、data/processed/summary.json、data/neo4j/*.csv、docs/neo4j_import.cypher、docs/cypher_queries.md、docs/deployment.md、docs/video_script.md。项目源代码包含 backend、frontend、scripts 三个主要目录。",
        ],
    ]


def main() -> None:
    summary = json.loads(SUMMARY.read_text(encoding="utf-8"))
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))

    blocks = report_sections(summary)
    placeholders = [p for p in doc.paragraphs if "具体内容***" in p.text]
    for paragraph, lines in zip(placeholders, blocks):
        add_block(paragraph, lines)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = "宋体"
            if run.font.size is None:
                run.font.size = Pt(10.5)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
